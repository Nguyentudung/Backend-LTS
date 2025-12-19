from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn


def _run_has_highlight(run) -> bool:
    """
    Determine whether a run is highlighted.

    Supports both the explicit highlight property and shaded runs that set the fill
    color (common when users "paint" text).
    """
    try:
        if run.font.highlight_color and run.font.highlight_color != WD_COLOR_INDEX.AUTO:
            return True
    except Exception:
        pass

    try:
        rpr = run._element.rPr  # noqa: SLF001
        if rpr is not None:
            shd = rpr.find(qn("w:shd"))
            if shd is not None:
                fill = (
                    shd.get(qn("w:fill"))
                    or shd.get("w:fill")
                    or shd.get("fill")
                    or ""
                )
                fill = fill.lower()
                # Word stores yellow as FFFF00; treat any non-empty fill as highlight.
                if fill and fill not in {"auto", "none", "000000"}:
                    return True
    except Exception:
        pass

    return False


def extract_text_with_highlight(docx_path: str):
    """
    Return list of paragraphs with highlight information:
    [
        {"text": "...", "has_highlight": True | False}
    ]
    """
    doc = Document(docx_path)
    results = []

    for p in doc.paragraphs:
        text = ""
        has_highlight = False

        for run in p.runs:
            text += run.text
            if _run_has_highlight(run):
                has_highlight = True

        if text.strip():
            results.append(
                {
                    "text": text,
                    "has_highlight": has_highlight,
                }
            )

    return results


def extract_paragraph_runs(docx_path: str):
    """
    Extract paragraphs with per-run highlight information for richer parsing.

    Returns:
    [
        {
            "text": "...",
            "has_highlight": True | False,
            "segments": [
                {"text": "...", "highlight": True | False},
                ...
            ],
        }
    ]
    """
    doc = Document(docx_path)
    results = []

    for p in doc.paragraphs:
        segments = []
        for run in p.runs:
            if not run.text:
                continue
            segments.append(
                {
                    "text": run.text,
                    "highlight": _run_has_highlight(run),
                }
            )

        text = "".join(seg["text"] for seg in segments)
        if text.strip():
            results.append(
                {
                    "text": text,
                    "has_highlight": any(seg["highlight"] for seg in segments),
                    "segments": segments,
                }
            )

    return results
